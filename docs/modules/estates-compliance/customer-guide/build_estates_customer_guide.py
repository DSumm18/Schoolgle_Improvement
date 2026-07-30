from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
SCREENSHOTS = ROOT / "screenshots"
OUTPUT = ROOT / "Schoolgle Estates Management Customer Guide.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_run(run, bold=False, color=None, size=None):
    run.bold = bold
    if color:
      run.font.color.rgb = RGBColor.from_string(color)
    if size:
      run.font.size = Pt(size)


def create_numbering_instance(doc, abstract_num_id):
    numbering = doc.part.numbering_part.element
    existing_ids = [
        int(num.get(qn("w:numId")))
        for num in numbering.findall(qn("w:num"))
    ]
    num_id = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    number_id = OxmlElement("w:numId")
    number_id.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(number_id)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string("2E74B5" if level < 3 else "1F4D78")
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.add_run(item)


def add_steps(doc, items):
    num_id = create_numbering_instance(doc, 7)
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        apply_numbering(paragraph, num_id)
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.add_run(item)


def add_callout(doc, title, body, fill="E8EEF5"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.05)
    p.paragraph_format.right_indent = Inches(0.05)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "8")
        border.set(qn("w:color"), "D0D7DE")
        borders.append(border)
    p_pr.append(borders)
    title_run = p.add_run(title)
    style_run(title_run, bold=True, color="0B2545")
    p.add_run(f"\n{body}")


def add_image(doc, filename, caption, width=6.5):
    path = SCREENSHOTS / filename
    if path.exists():
        shape = doc.add_picture(str(path), width=Inches(width))
        shape._inline.docPr.set("descr", caption)
        shape._inline.docPr.set("title", caption)
        doc.paragraphs[-1].paragraph_format.keep_with_next = True
        caption_p = doc.add_paragraph(caption)
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.paragraph_format.space_after = Pt(8)
        caption_p.runs[0].font.size = Pt(9)
        caption_p.runs[0].font.color.rgb = RGBColor.from_string("555555")


def build():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Schoolgle Estates Management Customer Guide")
    style_run(run, bold=True, color="0B2545", size=22)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Assets · Contractors · Compliance Checks · Tickets · Evidence").italic = True

    add_callout(
        doc,
        "Purpose",
        "This guide explains how a school uses Estates Management to keep assets, contractors, inspections, tickets, tasks, warranties and evidence connected in one audit trail.",
    )

    add_heading(doc, "1. The Full Circle", 1)
    doc.add_paragraph(
        "The Estates Management app is designed around one joined-up operational loop: add the asset, link it to the contractor and compliance check, complete the inspection, raise tickets or tasks when something fails, then keep evidence against the asset and check history."
    )
    add_bullets(
        doc,
        [
            "Assets are the real items on site: outlets, showers, extinguishers, boilers, emergency lights, playground items and similar equipment.",
            "Contractors are suppliers, maintainers, inspectors, warranty providers, or repair companies.",
            "Compliance checks are statutory, good-practice, or school-created inspections.",
            "Tickets and tasks capture follow-up work when an inspection finds an issue.",
            "Evidence records photos, certificates, completion notes, invoices, receipts and warranty documents.",
        ],
    )
    add_image(doc, "01-compliance-check-full-circle.png", "Compliance check with linked assets, tickets, tasks and completion history.")

    add_heading(doc, "2. Add an Asset", 1)
    add_steps(
        doc,
        [
            "Open Estates → Assets or Asset Tags.",
            "Select Add Asset.",
            "Enter the asset name, type, code, location and status.",
            "Add purchase value, invoice, purchase order, warranty and expected life details.",
            "Select the supplier and the maintainer/warranty contractor where known.",
            "Link the relevant compliance domains and checks, then save.",
            "Print the QR code and attach it to the item if required.",
        ],
    )
    add_image(doc, "03-add-asset-form.png", "Add Asset form showing purchase, warranty and contractor fields.")

    add_heading(doc, "3. Add a Contractor", 1)
    add_steps(
        doc,
        [
            "Open Estates → Contractors.",
            "Select Add Contractor.",
            "Enter the company, contact details, services, accreditation notes and status.",
            "Mark the contractor as preferred if appropriate.",
            "Save the contractor, then link it to assets, tasks, tickets or contracts.",
        ],
    )
    add_image(doc, "04-add-contractor-form.png", "Add Contractor form for supplier, maintainer and inspector records.")

    add_heading(doc, "4. Link Assets and Contractors to Checks", 1)
    doc.add_paragraph(
        "Use the asset record to decide which compliance checks apply to that item. For example, a shower can be linked to weekly outlet flushing, while an extinguisher can be linked to fire safety checks. Once linked, the check page shows the asset under Assets covered."
    )
    add_callout(
        doc,
        "Good practice",
        "Do not delete old assets when they are replaced if their history is needed. Retire or dispose of the old asset, add the new one, then print a new QR label for the replacement.",
        fill="F4F6F9",
    )

    add_heading(doc, "5. Complete a Compliance Check", 1)
    add_steps(
        doc,
        [
            "Open Estates → Compliance Checks.",
            "Choose the domain, such as Legionella Control or Fire Safety.",
            "Open the check.",
            "Review frequency, statutory/advisory status, next due date and linked assets.",
            "Expand Record completion.",
            "Enter status, inspection date, notes and evidence.",
            "Save the completion and review the Completion history section.",
        ],
    )

    add_callout(
        doc,
        "How to prove the record is stored",
        "Return to the domain list and confirm the check shows Completed. Reopen it, verify the next due date, notes and evidence in Completion history, open the evidence file, then reload the page and confirm the same information remains.",
        fill="E8F5F1",
    )
    doc.add_paragraph(
        "The system calculates the next due date from the recurrence. In the Rawdon acceptance test, the private non-statutory DEMO - Monthly Site Security Walkround was completed on 30 July 2026 and automatically moved to 30 August 2026."
    )
    add_image(doc, "11-rawdon-demo-check-in-list.png", "The completed monthly demo check is visible in the Security domain list.")
    add_image(doc, "08-rawdon-demo-check-completed.png", "The reopened check shows Fully Completed and the next due date of 30 August 2026.")
    add_image(doc, "09-rawdon-demo-history-and-image.png", "Completion history retains the date, person, notes, next due date and linked image.")
    add_image(doc, "10-rawdon-demo-evidence-open.png", "The saved image opens from the completion history.")
    add_callout(
        doc,
        "Customer-data guardrail",
        "Use a clearly labelled private demo check for demonstrations. Do not mark a real statutory check complete unless the school has genuinely completed it and the evidence is accurate. Use Aurora Primary to demonstrate invented failures and automatic tickets.",
        fill="FFF2CC",
    )

    add_heading(doc, "6. Raise a Ticket or Task from a Failed Check", 1)
    add_steps(
        doc,
        [
            "Open the compliance check.",
            "Select Raise ticket for faults, hazards, missing evidence or non-compliance.",
            "Select Add task for planned follow-up work.",
            "Choose the affected asset where relevant.",
            "Attach photos or documents.",
            "Assign staff or a contractor and save.",
        ],
    )
    doc.add_paragraph(
        "The ticket or task remains linked to the compliance check, the asset, the contractor and the helpdesk/task list."
    )
    add_image(doc, "05-raise-ticket-from-check.png", "New ticket opened from a compliance check with the check context preserved.")

    add_heading(doc, "7. Create a School Check", 1)
    add_steps(
        doc,
        [
            "Open the relevant domain, such as Fire Safety.",
            "Select Add Check.",
            "Start from a template or from scratch.",
            "Add name, description, evidence requirements, checklist items and frequency.",
            "Choose School check / non-statutory for local routines.",
            "Choose Statutory / regulated only where a law, regulation, trust policy or approved Schoolgle strategy applies.",
            "Add the statutory reference if the check is statutory, then save.",
        ],
    )
    add_image(
        doc,
        "06-custom-check-wizard.png",
        "Custom check wizard with statutory/non-statutory classification.",
        width=5.6,
    )
    add_callout(
        doc,
        "Statutory frequency rule",
        "Statutory frequencies are not school preferences. They should come from the regulation, approved strategy, or Schoolgle-managed statutory check library. School-created statutory checks are frequency-locked once saved.",
        fill="FFF2CC",
    )

    add_heading(doc, "8. Customer Testing Checklist", 1)
    add_bullets(
        doc,
        [
            "Add an asset with purchase, warranty and contractor information.",
            "Add a contractor and mark whether it is preferred.",
            "Link an asset to a compliance check.",
            "Complete a compliance check and review completion history.",
            "Raise a ticket from a failed check and confirm it appears on the check page.",
            "Raise a task from a check and assign it to a staff member or contractor.",
            "Replace an asset by retiring the old record and adding a new QR-tagged asset.",
            "Create a non-statutory school check and confirm it appears in the relevant domain.",
            "Complete a clearly labelled monthly demo check, then confirm the next due date advances by one month.",
            "Leave and reopen the check, then confirm the notes and evidence remain in Completion history.",
            "Open the saved image from the completion record.",
        ],
    )

    footer = section.footer.paragraphs[0]
    footer.text = "Schoolgle Estates Management Customer Guide"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor.from_string("555555")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
